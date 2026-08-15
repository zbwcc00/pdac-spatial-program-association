"""Build the v7 manuscript with author-annotated single-cell attribution QC."""
from pathlib import Path
import importlib.util
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
spec = importlib.util.spec_from_file_location("docx_v5_for_v6", ROOT / "scripts/build_manuscript_docx_v5.py")
module = importlib.util.module_from_spec(spec)
spec.loader.exec_module(module)
module.SOURCE = ROOT / "manuscript/drafts/Manuscript_full_v7_single_cell_attribution.md"
module.OUTPUT = ROOT / "manuscript/drafts/Manuscript_full_v7_single_cell_attribution.docx"
module.MAIN = ROOT / "figures/main_v4"
def figures_v6(doc, legends):
    doc.add_page_break()
    paragraph = doc.add_paragraph(style="Heading 1")
    module.inline(paragraph, "Main figures", 16, module.BLUE)
    paths = [
        module.MAIN / "Figure1_v4_primary_and_spatial_nulls.png",
        ROOT / "figures/submission_v6/Figure2_v6_GSE278687_representative_spatial_maps.png",
        module.MAIN / "Figure2_v4_robustness_and_competition.png",
        module.MAIN / "Figure3_v4_external_technical_replication.png",
        ROOT / "figures/main_v2/Figure5_v1_spatial_holdout_validation.png",
    ]
    for path, legend in zip(paths, legends):
        paragraph = doc.add_paragraph()
        paragraph.alignment = module.WD_ALIGN_PARAGRAPH.CENTER
        shape = paragraph.add_run().add_picture(str(path), width=module.Inches(6.35))
        shape._inline.docPr.set("descr", legend)
        shape._inline.docPr.set("title", legend.split(".", 1)[0])
        paragraph = doc.add_paragraph(style="Caption")
        module.inline(paragraph, legend, 9.5, module.MUTED)
module.figures = figures_v6
module.re = __import__("re")
module.build()

# The v5 builder is retained for its established layout, but its main tables
# were hard-coded for the prior global-Moran-only manuscript. Replace only
# the inherited rows that are version-sensitive after the initial build.
doc = Document(module.OUTPUT)
for section in doc.sections:
    header = section.header.paragraphs[0]
    header.text = "PDAC local score-field association | v7 author-review manuscript"
    header.alignment = module.WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        module.font(run, 8.5, module.MUTED, italic=True)
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("v5 manuscript with extended spatial-null inference"):
        paragraph.text = "v7 manuscript with spatial sensitivities and author-annotated single-cell attribution QC"
        for run in paragraph.runs:
            module.font(run, 10.5, module.MUTED, italic=True)
    if paragraph.text.startswith(("A reproducible local spatial association of mregDC-like", "Local association of mregDC-like and T-cell/lymphoid transcriptional programs in PDAC under")):
        paragraph.text = "Local association of mregDC-like and T-cell/lymphoid transcriptional programs in pancreatic ductal adenocarcinoma"
        for run in paragraph.runs:
            module.font(run, 20, module.BLUE, bold=True)
    if paragraph.text == "PDAC local immune-program organization | v5 author-review manuscript":
        paragraph.text = "PDAC local score-field association | v7 author-review manuscript"
        for run in paragraph.runs:
            module.font(run, 8.5, module.MUTED, italic=True)

table = doc.tables[2]
headers = ["Check", "GSE278687", "GSE277116", "Methodological boundary"]
rows = [
    ["Array-block null (principal)", "999 draws; 0/999; add-one P=0.001; 95% null 0.24605-0.29417", "Not applied", "Patient-level hierarchy-matched randomization; primary evidence"],
    ["Global-Moran graph sensitivity", "1,000; 0/1,000; add-one Monte Carlo P=0.000999; global mismatch 2.1385e-05; mask audit median 0.03518, max 0.05497", "1,000; 0/1,000; add-one Monte Carlo P=0.000999; global mismatch 1.6161e-05; mask audit median 0.02733, max 0.04810", "Matches whole-section Moran's I; mask quantity is an audit"],
    ["Mask-Moran graph sensitivity", "1,000; 0/1,000; add-one Monte Carlo P=0.000999; mask mismatch 2.0851e-05; global audit median 0.03523, max 0.05810", "1,000; 0/1,000; add-one Monte Carlo P=0.000999; mask mismatch 1.9585e-05; global audit median 0.02630, max 0.04315", "Matches DC-core-mask Moran's I; global quantity is an audit"],
    ["Primary-score zero fractions", "mregDC 0.277; Tfh 0.140", "mregDC 0.538; Tfh 0.529", "Spot-level program scores, not cell abundance"],
]
for row, values in zip(table.rows, [headers, *rows]):
    for cell, value in zip(row.cells, values):
        cell.text = value
        module.margins(cell)
        cell.vertical_alignment = module.WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            module.font(run, 7.6 if row is not table.rows[0] else 7.8, module.DARK if row is table.rows[0] else module.INK, bold=(row is table.rows[0]))
module.mark_header(table.rows[0])
module.table_geometry(table, [1820, 2520, 2520, 2150])

# Add the post-hoc sensitivity table after the inherited main tables. This
# makes the conditional estimand and its tested alternatives visible in the
# author-review DOCX rather than leaving them only in the supplement.
module.add_table(doc, "Table 4. Post-hoc conditioning and neighbourhood-scale sensitivity", ["Setting", "GSE278687 (18 patients)", "GSE277116 (18 samples)", "Status"], [
    ["Whole tissue, no DC-core adjustment", "0.399 (0.327-0.496)", "0.383 (0.281-0.475)", "Post-hoc conditional-estimand sensitivity"],
    ["Whole tissue, DC-core adjustment", "0.376 (0.277-0.452)", "0.369 (0.252-0.439)", "Post-hoc conditional-estimand sensitivity"],
    ["DC-core > median, adjusted", "0.428 (0.313-0.523)", "0.413 (0.293-0.490)", "Post-hoc mask sensitivity"],
    ["DC-core > Q75, adjusted", "0.449 (0.360-0.551)", "0.419 (0.322-0.532)", "Post-hoc mask sensitivity"],
    ["k=4 / 6 / 12, adjusted mean mask", "0.375 / 0.418 / 0.475", "0.372 / 0.414 / 0.473", "Post-hoc local-scale sensitivity"],
], [2300, 2400, 2400, 2350])

module.add_table(doc, "Table 5. Independent author-annotated single-cell program-attribution QC (GSE202051)", ["Comparison", "Patient-level result", "Interpretation", "Boundary"], [
    ["mregDC-like: activated DC vs cDC2 (sole designated primary test)", "23 paired patients; median difference 0.479 (95% CI 0.424-0.668); two-sided P=2.38e-6", "Fixed mregDC-like score is higher in author-labelled activated DCs", "Program attribution only; not a spatial or cell-interaction test; not prospectively preregistered"],
    ["mregDC-like minimum-cell sensitivities", ">=3 cells/label: 9 patients, difference 0.718 (95% CI 0.359-0.933), two-sided P=0.00391, FDR q=0.00641; >=5: 7 patients, difference 0.718 (95% CI 0.359-0.747), P=0.0156, q=0.0169", "Positive under stricter patient-level cell-count requirements", "Exploratory; reduced retained patient count"],
    ["mregDC-like minus CCL19: activated DC vs cDC2", "23 paired patients; median difference 0.573; exploratory FDR q=1.40e-5", "The contrast persists after CCL19 deletion", "Tests dependence of the predefined score; does not establish molecular contribution or mechanism; CCL19 detected in 1/91 activated DCs"],
    ["Tfh-like: CD4+ T vs Treg", "25 paired patients; median difference 0.061; exploratory FDR q=0.00770", "Small broad T-cell-state contrast", "No author-defined Tfh category; not Tfh-specific"],
    ["Tfh-like minus IL7R: CD4+ T vs Treg", "Median difference -0.083", "Direction reverses after IL7R deletion", "Exploratory sensitivity; Tfh identity cannot be claimed"],
], [2450, 2500, 2350, 2150])
doc.save(module.OUTPUT)
