"""Build v5 author-review DOCX from the versioned spatial-controls manuscript."""
from __future__ import annotations
from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript/drafts/Manuscript_full_v5_spatial_controls.md"
OUTPUT = ROOT / "manuscript/drafts/Manuscript_full_v5_spatial_controls.docx"
MAIN = ROOT / "figures/main_v3"
HOLDOUT = ROOT / "figures/main_v2/Figure5_v1_spatial_holdout_validation.png"
BLUE, DARK, INK, MUTED = RGBColor(46,116,181), RGBColor(31,77,120), RGBColor(20,20,20), RGBColor(90,90,90)
WIDTH = 9360

def font(run, size=11, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name); run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.font.color.rgb = color
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def margins(cell):
    p = cell._tc.get_or_add_tcPr(); node = p.first_child_found_in("w:tcMar")
    if node is None: node = OxmlElement("w:tcMar"); p.append(node)
    for side, val in (("top",80),("start",120),("bottom",80),("end",120)):
        e = node.find(qn(f"w:{side}"))
        if e is None: e = OxmlElement(f"w:{side}"); node.append(e)
        e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa")

def shade(cell, fill="E8EEF5"):
    e = OxmlElement("w:shd"); e.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(e)

def mark_header(row):
    props = row._tr.get_or_add_trPr(); e = OxmlElement("w:tblHeader"); e.set(qn("w:val"), "true"); props.append(e)

def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    p = table._tbl.tblPr
    for tag, value in (("w:tblW",sum(widths)),("w:tblInd",120)):
        e = p.first_child_found_in(tag)
        if e is None: e = OxmlElement(tag); p.append(e)
        e.set(qn("w:w"),str(value)); e.set(qn("w:type"),"dxa")
    grid = table._tbl.tblGrid
    for col in list(grid): grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"),str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell._tc.tcPr.tcW.set(qn("w:w"),str(widths[i])); cell._tc.tcPr.tcW.set(qn("w:type"),"dxa")

def inline(p, text, size=11, color=INK):
    pos = 0
    for m in re.finditer(r"(\*\*.+?\*\*|`.+?`)", text):
        if m.start() > pos: font(p.add_run(text[pos:m.start()]), size, color)
        token=m.group(0)
        if token.startswith("**"): font(p.add_run(token[2:-2]), size, color, bold=True)
        else: font(p.add_run(token[1:-1]), max(9,size-.5), DARK, name="Consolas")
        pos=m.end()
    if pos < len(text): font(p.add_run(text[pos:]), size, color)

def configure(doc):
    sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(1); sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.49)
    normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal.font.color.rgb=INK; normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.25
    for name,size,before,after in (("Heading 1",16,18,10),("Heading 2",13,14,7)):
        s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=BLUE; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    cap=doc.styles["Caption"]; cap.font.name="Calibri"; cap.font.size=Pt(9.5); cap.font.italic=True; cap.font.color.rgb=MUTED; cap.paragraph_format.space_before=Pt(4); cap.paragraph_format.space_after=Pt(10)
    p=sec.header.paragraphs[0]; p.text=""; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p=sec.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("Page "),9,MUTED)
    run=p.add_run(); b=OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"),"begin"); i=OxmlElement("w:instrText"); i.set(qn("xml:space"),"preserve"); i.text=" PAGE "; e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end"); run._r.extend([b,i,e]); font(run,9,MUTED)

def title(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(6)
    font(p.add_run("Local association of mregDC-like and T-cell/lymphoid transcriptional programs in PDAC under specified spatial scoring definitions"),20,RGBColor(24,52,77),bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(13); font(p.add_run("v5 manuscript with extended spatial-null inference and scoreability audit"),10.5,MUTED,italic=True)
    for lab,val in (
        ("Authors", "Bowen Zheng, Yuqiao Sun, Jinyuan Chi, Hao Li*"),
        ("Affiliations", "Division of Hepatobiliary Pancreatic Surgery, The Affiliated Hospital of Yanbian University, Yanji 133000, China"),
        ("Corresponding author", "Hao Li, lih@ybu.edu.cn"),
    ):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); font(p.add_run(lab+": "),10,DARK,bold=True); font(p.add_run(val),10,MUTED)

def add_table(doc, label, heads, rows, widths):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4); font(p.add_run(label),10.5,DARK,bold=True)
    t=doc.add_table(rows=1,cols=len(heads)); t.style="Table Grid"
    for j, head in enumerate(heads):
        c=t.rows[0].cells[j]; c.text=head; c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT; shade(c); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: font(r,8.4,DARK,bold=True)
    mark_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for j,value in enumerate(row):
            cells[j].text=str(value); cells[j].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT; margins(cells[j]); cells[j].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for r in cells[j].paragraphs[0].runs: font(r,8.2,INK)
    table_geometry(t,widths)

def figures(doc, legends):
    doc.add_page_break(); p=doc.add_paragraph(style="Heading 1"); inline(p,"Main figures",16,BLUE)
    paths=[MAIN/"Figure1_v3_primary_and_spatial_nulls.png", MAIN/"Figure2_v3_robustness_and_competition.png", MAIN/"Figure3_v3_external_technical_replication.png", HOLDOUT]
    for path,legend in zip(paths,legends):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
        shape=p.add_run().add_picture(str(path),width=Inches(6.35)); shape._inline.docPr.set("descr", legend); shape._inline.docPr.set("title", legend.split(".", 1)[0])
        p=doc.add_paragraph(style="Caption"); inline(p,legend,9.5,MUTED)

def build():
    doc=Document(); configure(doc); title(doc)
    source=SOURCE.read_text(encoding="utf-8"); body, legends=source.split("## Figure legends",1)
    seen=False
    for raw in body.splitlines():
        line=raw.strip()
        if not line: continue
        if line.startswith("# "):
            if seen: continue
            seen=True; continue
        if line.startswith("**Running title:") or line.startswith("**Authors:") or line.startswith("**Affiliations:") or line.startswith("**Corresponding author:"): continue
        if line.startswith("## "): p=doc.add_paragraph(style="Heading 1"); inline(p,line[3:],16,BLUE); continue
        if line.startswith("### "): p=doc.add_paragraph(style="Heading 2"); inline(p,line[4:],13,BLUE); continue
        p=doc.add_paragraph(); inline(p,line,11,INK)
    doc.add_page_break(); p=doc.add_paragraph(style="Heading 1"); inline(p,"Main tables",16,BLUE)
    add_table(doc,"Table 1. Cohorts and evidentiary roles",["Dataset","Platform","Units","Role","Inference boundary"],[
        ["GSE278687","Visium","21 sections / 18 patients","Primary spatial cohort","Section medians collapsed to patients"],
        ["GSE277116","Spatial packages","18 tumour packages","External technical replication","No patient IDs; sample-level only"],
        ["GSE217845","scRNA-seq","10 tumours","Supplementary QC","Marker gating is circular, not validation"],
        ["GSE237846/GSE205354","Targeted / matrices","Discovery context","Descriptive context","Not independent spatial inference"],
    ],[1250,1300,1500,2050,3260])
    add_table(doc,"Table 2. Primary and competitive-program results",["Endpoint","GSE278687 (18 patients)","GSE277116 (18 samples)","Interpretation"],[
        ["Primary local r","0.418 (0.312-0.516)","0.414 (0.287-0.470)","Local program association"],
        ["After non-overlap broad-T adjustment","0.331 (0.175-0.396)","0.350 (0.242-0.396)","Residual, not Tfh specificity"],
        ["mregDC score without CCL19","0.303 (0.183-0.396)","0.203 (0.117-0.315)","Part of this predefined-score association depends on CCL19"],
        ["Joint-model Tfh beta","0.158 (0.089-0.205)","0.169 (0.124-0.217)","Specified-model residual"],
        ["Non-Tfh CD4 local r","0.436","0.440","Limits Tfh-specific interpretation"],
        ["Macrophage local r","-0.020","0.009","Negative competitive reference"],
    ],[2250,2050,2050,3010])
    add_table(doc,"Table 3. Spatial-inference and scoreability audit",["Check","GSE278687","GSE277116","Boundary"],[
        ["Corrected block null","999 draws; P=0.001; range 0.246-0.294","Not applied","Patient-level hierarchy matched"],
        ["Graph null","1,000 draws; P=0.001; global Moran mismatch 2.14e-05","1,000 draws; P=0.001; mismatch 1.62e-05","Calibrated globally, not in mask"],
        ["Mask Moran audit","Median absolute mismatch 0.035; max 0.055","Median 0.027; max 0.048","Audit only, not calibration"],
        ["Primary-score zero fractions","mregDC 0.277; Tfh 0.140","mregDC 0.538; Tfh 0.529","Spot-level programs, not abundance"],
    ],[2150,2400,2400,2410])
    blocks=re.findall(r"### Figure \d+\..*?(?=\n### Figure |\Z)",legends,flags=re.S)
    clean=[re.sub(r"^###\s*", "", re.sub(r"\s+"," ",x.strip())) for x in blocks[:5]]
    figures(doc,clean)
    doc.save(OUTPUT); print(OUTPUT)

if __name__ == "__main__": build()
